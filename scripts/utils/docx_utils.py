"""Word document (.docx) manipulation utilities.

Handles reading, writing, bookmark insertion, and hyperlink creation.
"""

import re
import shutil
from datetime import datetime
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def open_docx(path: str) -> Document:
    """Open a .docx file and return Document object."""
    return Document(path)


def backup_docx(path: str) -> str:
    """Create a timestamped backup of the docx file. Returns backup path."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = path.replace('.docx', f'_backup_refs_{timestamp}.docx')
    shutil.copy2(path, backup_path)
    return backup_path


def extract_all_text(doc: Document) -> list:
    """Extract all paragraph texts as a list of strings."""
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def extract_reference_section(doc: Document) -> tuple:
    """Extract the reference list section.

    Returns:
        (ref_title_para_index, ref_entries) where ref_entries is a list of
        (paragraph_index, paragraph_object, ref_number)
    """
    ref_title_idx = None
    ref_entries = []
    non_ref_count = 0
    max_blanks = 3

    section_headers = (
        '参考文献', 'References', 'REFERENCES', 'BIBLIOGRAPHY',
        '附录', 'Appendix', '致谢', 'Acknowledgments', 'Acknowledgements',
    )

    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()

        if ref_title_idx is None:
            if t in ('参考文献', 'References', 'REFERENCES', 'BIBLIOGRAPHY'):
                ref_title_idx = i
            continue

        # Stop if we hit another section
        if t in ('附录', 'Appendix', '致谢', 'Acknowledgments', 'Acknowledgements',
                 'Abstract', '摘要', 'Introduction', '引言', 'Conclusion', '结论',
                 '图表索引', 'List of Figures', 'List of Tables', 'Appendix'):
            break

        if re.match(r'^\[\d+\]', t):
            non_ref_count = 0
            num = int(re.match(r'\[(\d+)\]', t).group(1))
            ref_entries.append((i, para, num))
        elif not t:
            non_ref_count += 1
            if non_ref_count > max_blanks:
                break
        else:
            non_ref_count += 1
            if non_ref_count > 2:
                break

    return ref_title_idx, ref_entries


def extract_citations_from_text(doc: Document) -> list:
    """Extract all citation markers from body text (not reference section).

    Supports [N], [N-M], [N,M], [N,M-O] formats.
    Returns list of citation numbers in order of first appearance.
    """
    first_occurrence = []
    seen_nums = set()

    ref_end = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() in ('参考文献', 'References', 'REFERENCES', 'BIBLIOGRAPHY'):
            ref_end = i
            break

    for i, para in enumerate(doc.paragraphs):
        if ref_end is not None and i >= ref_end:
            break
        text = para.text.strip()
        if not text:
            continue
        for match in re.finditer(r'\[([0-9,\-\s]+)\]', text):
            content = match.group(1)
            for num in _parse_citation_range(content):
                if num not in seen_nums:
                    seen_nums.add(num)
                    first_occurrence.append(num)

    return first_occurrence


def _parse_citation_range(content: str) -> list:
    """Parse citation content like '1-3,5,7-9' into flat list of ints."""
    nums = []
    for part in content.split(','):
        part = part.strip()
        if '-' in part:
            a, _, b = part.partition('-')
            try:
                for n in range(int(a.strip()), int(b.strip()) + 1):
                    nums.append(n)
            except ValueError:
                pass
        else:
            try:
                nums.append(int(part.strip()))
            except ValueError:
                pass
    return nums


def extract_citation_context(doc: Document, num: int) -> str:
    """Find the context around a citation marker [num] in body text."""
    for para in doc.paragraphs:
        text = para.text
        pattern = f'[{num}]'
        if pattern in text:
            idx = text.find(pattern)
            start = max(0, idx - 20)
            end = min(len(text), idx + len(pattern) + 30)
            return text[start:end]
    return ''


def add_bookmark_to_paragraph(para, bookmark_name: str) -> bool:
    """Add a bookmark at the start of a paragraph."""
    try:
        bookmark_id = str(abs(hash(bookmark_name)) % 1000000)
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), bookmark_id)
        bookmark_start.set(qn('w:name'), bookmark_name)

        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), bookmark_id)

        first_run = para.runs[0] if para.runs else None
        if first_run:
            first_run._r.addprevious(bookmark_start)
            para._p.append(bookmark_end)
            return True
        return False
    except Exception:
        return False


def make_hyperlink_element(ref_text: str, bookmark_name: str) -> OxmlElement:
    """Create a w:hyperlink XML element pointing to a bookmark."""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = ref_text
    r.append(t)
    hyperlink.append(r)
    return hyperlink


def replace_citation_with_hyperlink(para, ref_num: int) -> bool:
    """Replace [ref_num] citation in a paragraph with a hyperlink.

    Returns True if replacement succeeded.
    """
    text = para.text
    ref_text = f'[{ref_num}]'
    start = text.find(ref_text)
    if start < 0:
        return False

    bookmark_name = f"Ref_{ref_num}"

    # Find the run containing this citation
    char_count = 0
    target_run = None
    run_offset = -1

    for run in para.runs:
        run_len = len(run.text)
        if char_count <= start < char_count + run_len:
            target_run = run
            run_offset = start - char_count
            break
        char_count += run_len

    if target_run is None:
        return False

    # Save format properties
    rPr = None
    for child in target_run._r:
        if child.tag == qn('w:rPr'):
            rPr = deepcopy(child)
            break

    before = target_run.text[:run_offset]
    after = target_run.text[run_offset + len(ref_text):]

    hyperlink_elem = make_hyperlink_element(ref_text, bookmark_name)

    if before:
        target_run.text = before
        target_run._r.addnext(hyperlink_elem)
    else:
        target_run.text = ''
        ref = target_run._r
        parent = ref.getparent() if hasattr(ref, 'getparent') else ref
        ref.addprevious(hyperlink_elem)

    if after:
        new_run = OxmlElement('w:r')
        if rPr:
            new_run.append(deepcopy(rPr))
        new_t = OxmlElement('w:t')
        new_t.text = after
        new_run.append(new_t)
        hyperlink_elem.addnext(new_run)

    return True


def insert_text_at_run(para, search_text: str, insert_text: str) -> bool:
    """Insert text after the first occurrence of search_text in any run."""
    for run in para.runs:
        if search_text in run.text:
            idx = run.text.index(search_text) + len(search_text)
            run.text = run.text[:idx] + insert_text + run.text[idx:]
            return True
    return False


def get_reference_text(ref_entries: list, ref_num: int) -> str:
    """Get the full text of a reference entry by its number."""
    for idx, para, num in ref_entries:
        if num == ref_num:
            return para.text.strip()
    return ''


def update_reference_text(para, new_text: str) -> bool:
    """Update the text of a reference paragraph."""
    if not para.runs:
        return False
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ''
    return True
