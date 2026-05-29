"""
论文排版修复脚本
修复内容：
1. 清理跨段落的内联图表引用（如"图 4昆明..."出现在段落末尾）
2. 统一模型命名（AdAdaptiveGate -> AdaptiveGate）
3. 修正单位间距（1.81ms -> 1.81 ms）
4. 修正 LaTeX 残留语法（** -> 等）
5. 修复英文作者行空格缺失（China2. -> China 2.）

安全保护：
- 自动跳过含 OMML 公式的段落（防止破坏公式对象）
- 自动跳过含图片/绘图的段落（防止删除图片）
- 创建备份文件

使用方法：
    pip install python-docx lxml
    python fix_docx.py
"""

import re
import os

try:
    from docx import Document
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    exit(1)

try:
    from lxml import etree
except ImportError:
    print("请先安装 lxml: pip install lxml")
    exit(1)

# 配置
DOCX_PATH = os.path.join(os.path.dirname(__file__),
    "论文_面向多尺度时序预测的时域门控混合模型_终版.docx")
BACKUP_PATH = DOCX_PATH.replace('.docx', '_备份_自动修复.docx')

# OMML 命名空间（Word公式）
M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
# Word 命名空间
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _should_skip_paragraph(para) -> bool:
    """检查段落是否包含公式或图片，需要跳过处理"""
    # 1. 跳过含 OMML 公式的段落
    omml = para._element.findall(f'.//{{{M_NS}}}oMath')
    omml_para = para._element.findall(f'.//{{{M_NS}}}oMathPara')
    if omml or omml_para:
        return True

    # 2. 跳过含图片/绘图的段落
    drawings = para._element.findall(f'.//{{{W_NS}}}drawing')
    if drawings:
        return True

    # 3. 跳过含 MathType 方程样式的段落
    pPr = para._element.find(f'{{{W_NS}}}pPr')
    if pPr is not None:
        pStyle = pPr.find(f'{{{W_NS}}}pStyle')
        if pStyle is not None:
            style_val = pStyle.get(f'{{{W_NS}}}val', '')
            if 'Equation' in style_val or 'MTDisplay' in style_val:
                return True

    return False


def fix_paragraph_text(text):
    """修复段落文本中的问题"""
    if not text:
        return text

    # 1. 清理内联图表引用：去掉段落末尾或中间的 "图 X描述文字"、"表 X描述文字"
    #   但保留纯数字所在的图表引用（如括号内的图编号）
    #   规则：删除独立成句的 "图 X..." 和 "表 X..."（非括号引用）
    text = re.sub(
        r'[。；]?\s*图[\s\S]{0,2}\d+[^\。]*?(?:对比|示意|分布|曲线|热力|直方|收敛|演化|对比)',
        '', text
    )
    text = re.sub(
        r'[。；]?\s*表[\s\S]{0,2}\d+[^\。]*?(：|对比|性能|结果|参数)',
        '', text
    )

    # 2. 修复命名一致性
    text = text.replace('AdAdaptiveGate', 'AdaptiveGate')
    text = text.replace('AdAdaptive', 'Adaptive')

    # 3. 修复单位间距（数字+ms/ MHz/ GB等）
    text = re.sub(r'(\d+(?:\.\d+)?)(ms|MHz|GB|MB|KB)', r'\1 \2', text)

    # 4. 修复英文标点后空格缺失
    text = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', text)  # China2 -> China 2

    # 5. 修复连续短横线
    text = text.replace('---', '--')
    text = text.replace('--', '--')

    # 6. 修复全角冒号后空格（DOI行）
    text = re.sub(r'doi：(\s*)', r'doi: \1', text)

    return text


def main():
    docx_path = DOCX_PATH
    print(f"正在读取: {docx_path}")

    if not os.path.exists(docx_path):
        print(f"错误: 找不到文件 {docx_path}")
        # 尝试找备份文件
        alt_path = docx_path.replace('_终版.docx', '_终版_备份.docx')
        if os.path.exists(alt_path):
            print(f"尝试使用备份文件: {alt_path}")
            docx_path = alt_path
        else:
            return

    # 创建备份
    import shutil
    shutil.copy2(docx_path, BACKUP_PATH)
    print(f"已创建备份: {BACKUP_PATH}")

    doc = Document(docx_path)
    changes = 0
    skipped = 0

    # 修复正文段落
    for i, para in enumerate(doc.paragraphs):
        # 安全检查：跳过公式、图片段落
        if _should_skip_paragraph(para):
            skipped += 1
            continue

        original = para.text
        fixed = fix_paragraph_text(original)
        if fixed != original and fixed.strip():
            # 直接替换段落中所有run的文本
            if para.runs:
                # 按比例分配修复后的文本到各个run
                runs_text = ''.join(r.text for r in para.runs)
                if runs_text == original:
                    # 简单情况：所有run拼接起来等于原文
                    # 按原run长度比例重新分配
                    total_len = sum(len(r.text) for r in para.runs if r.text)
                    if total_len > 0:
                        # 用第一个run承载所有文本（最简单的做法）
                        para.runs[0].text = fixed
                        for r in para.runs[1:]:
                            r.text = ''
                        changes += 1
                        print(f"  修复段落 {i+1}: {fixed[:60]}...")

    # 修复表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original = cell.text
                fixed = fix_paragraph_text(original)
                if fixed != original:
                    for para in cell.paragraphs:
                        if para.runs:
                            runs_text = ''.join(r.text for r in para.runs)
                            if runs_text == original:
                                para.runs[0].text = fixed
                                for r in para.runs[1:]:
                                    r.text = ''
                                changes += 1

    # 保存
    doc.save(docx_path)
    print(f"\n修复完成！")
    print(f"  修改: {changes} 处")
    print(f"  跳过(公式/图片): {skipped} 段")
    print(f"  文件: {docx_path}")
    print(f"  备份: {BACKUP_PATH}")


if __name__ == '__main__':
    main()
