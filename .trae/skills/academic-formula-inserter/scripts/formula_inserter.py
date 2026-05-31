import sys
import os
import shutil
import argparse
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
sys.path.insert(0, str(Path(__file__).parent / "utils"))

from latex_omml_converter import latex_to_omml, detect_equations_in_docx, omml_to_latex
from formula_numbering import number_equations_in_docx

try:
    from utils.journal_profiles import get_profile, list_journals
except ImportError:
    from journal_profiles import get_profile, list_journals

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def cmd_insert_formula(args):
    if getattr(args, 'mathtype', False):
        _insert_mathtype_formula(args)
        return

    doc = Document(args.docx_path)

    omml = latex_to_omml(args.latex, display=args.display)

    insert_paragraph_index = args.position if args.position is not None else len(doc.paragraphs) - 1

    if insert_paragraph_index < 0:
        insert_paragraph_index = 0
    if insert_paragraph_index >= len(doc.paragraphs):
        insert_paragraph_index = len(doc.paragraphs) - 1

    para = doc.paragraphs[insert_paragraph_index]

    omml_element = parse_xml(omml)

    if args.display:
        omath_para = etree.SubElement(para._element, f"{{{M_NS}}}oMathPara")
        omath_para.append(omml_element)
    else:
        para._element.append(omml_element)

    if args.interactive:
        from formula_numbering import interactive_numbering_prompt, _style_id_to_format, _add_number_to_paragraph
        style = interactive_numbering_prompt("请选择编号格式")
        if style != "none":
            num_text = _style_id_to_format(style).format(num=args.number or 1)
            _add_number_to_paragraph(para, num_text)
            print(f"  Number: {num_text}")
    elif args.number is not None or args.style:
        from formula_numbering import _add_number_to_paragraph, _style_id_to_format
        style = args.style or "parentheses"
        profile = get_profile(style)
        if profile:
            num_text = profile["numbering_format"].format(num=args.number or 1)
        else:
            fmt = _style_id_to_format(style)
            num_text = fmt.format(num=args.number or 1)
        _add_number_to_paragraph(para, num_text)
        print(f"  Number: {num_text}")

    doc.save(args.docx_path)
    print(f"Formula inserted: {args.latex}")
    if args.display:
        print("  Type: display equation")


def _insert_mathtype_formula(args):
    """Insert formula using MathType-compatible format via ZIP-level injection."""
    import zipfile, io, re, tempfile
    from lxml import etree
    from mathtype_integration import (
        add_mt_display_equation_style,
        build_mt_omml_equation,
        inject_mt_paragraphs_via_zip,
    )

    temp_dir = Path(args.docx_path).parent
    temp_path = str(temp_dir / f"_mt_temp_{os.urandom(4).hex()}.docx")

    shutil.copy2(args.docx_path, temp_path)

    add_mt_display_equation_style(temp_path)

    eq_num = args.number or 1
    omml = latex_to_omml(args.latex, display=False)
    para_xml = build_mt_omml_equation(omml, eq_num)
    inject_mt_paragraphs_via_zip(temp_path, [para_xml])

    # Clean up python-docx corrupted elements if any exist
    with zipfile.ZipFile(temp_path, 'r') as zf:
        doc_xml = zf.read('word/document.xml')

    cleaned = re.sub(r'<m:sSupPr>.*?</m:sSupPr>', '', doc_xml.decode('utf-8'), flags=re.DOTALL)
    cleaned = re.sub(r'<m:sSubPr>.*?</m:sSubPr>', '', cleaned, flags=re.DOTALL)
    cleaned = re.sub(r'<m:ctrlPr>.*?</m:ctrlPr>', '', cleaned, flags=re.DOTALL)
    cleaned = re.sub(r'<m:sup/>', '', cleaned)
    cleaned = re.sub(r'<m:sub/>', '', cleaned)

    buffer = io.BytesIO()
    with zipfile.ZipFile(temp_path, 'r') as zf:
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as out:
            for item in zf.infolist():
                if item.filename == 'word/document.xml':
                    out.writestr(item, cleaned.encode('utf-8'))
                else:
                    out.writestr(item, zf.read(item.filename))

    shutil.move(temp_path, args.docx_path)
    print(f"MathType formula inserted: {args.latex}")
    print(f"  Type: {'display' if args.display else 'inline'} equation")
    print(f"  Number: {eq_num} (SEQ MTEqn)")
    print(f"  Style: MTDisplayEquation (double-click editable in MathType)")


def cmd_batch_convert(args):
    if args.backup:
        backup_path = args.docx_path.replace(".docx", "_backup.docx")
        shutil.copy2(args.docx_path, backup_path)
        print(f"Backup saved: {backup_path}")

    doc = Document(args.docx_path)

    latex_patterns = [
        (r'\$\$(.*?)\$\$', True),
        (r'\$(.*?)\$', False),
    ]

    converted = 0
    for para in doc.paragraphs:
        text = para.text
        for pattern, is_display in latex_patterns:
            import re
            matches = list(re.finditer(pattern, text, re.DOTALL))
            for match in reversed(matches):
                latex_code = match.group(1).strip()
                try:
                    omml = latex_to_omml(latex_code, display=is_display)
                    omml_element = parse_xml(omml)

                    if is_display:
                        omath_para = etree.SubElement(para._element, f"{{{M_NS}}}oMathPara")
                        omath_para.append(omml_element)
                    else:
                        para._element.append(omml_element)

                    range_start = match.start()
                    range_end = match.end()
                    text_before = text[:range_start]
                    text_after = text[range_end:]
                    para.clear()
                    run = para.add_run(text_before + text_after)
                    converted += 1
                except Exception as e:
                    print(f"  Warning: could not convert '{latex_code[:30]}...': {e}")

    doc.save(args.docx_path)
    print(f"Batch conversion complete. {converted} equations converted.")
    if args.backup:
        print(f"  Backup: {backup_path}")


def cmd_number_formulas(args):
    doc = Document(args.docx_path)

    profile = get_profile(args.style) if args.style and not args.interactive else None
    style = args.style or "parentheses"

    count = number_equations_in_docx(doc, style=style, profile=profile, interactive=args.interactive)
    doc.save(args.docx_path)
    if not args.interactive:
        print(f"Numbered {count} equations with style '{style}'")


def cmd_validate_format(args):
    doc = Document(args.docx_path)
    profile = get_profile(args.journal)

    if not profile:
        print(f"Unknown journal: {args.journal}")
        print("Available journals:")
        print(list_journals())
        return

    equations = detect_equations_in_docx(doc)
    issues = []

    for eq in equations:
        if not eq["is_display"] and profile.get("display_align") == "center":
            issues.append(f"  Eq {eq['index']} (para {eq['paragraph_index']}): "
                          f"inline equation in display-style journal")

    for para in doc.paragraphs:
        runs = para.runs
        for run in runs:
            if run.font.superscript:
                continue

    print(f"Validation for: {profile['name']}")
    print(f"  Font: {profile['font']}")
    print(f"  Numbering: {profile['numbering_format']}")
    print(f"  Equations detected: {len(equations)}")

    if issues:
        print(f"  Issues found: {len(issues)}")
        for issue in issues:
            print(issue)
    else:
        print("  No issues found. Format looks good!")


def cmd_export_latex(args):
    doc = Document(args.docx_path)
    equations = detect_equations_in_docx(doc)

    output_path = args.output or args.docx_path.replace(".docx", "_formulas.tex")

    lines = []
    lines.append("% Equations extracted from: " + os.path.basename(args.docx_path))
    lines.append(f"% Total: {len(equations)} equations")
    lines.append("")

    for eq in equations:
        sep = "$$" if eq["is_display"] else "$"
        lines.append(f"{sep}{eq['latex']}{sep}")
        lines.append("")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    print(f"Exported {len(equations)} equations to: {output_path}")


def cmd_list_styles(args):
    print(list_journals())


def main():
    parser = argparse.ArgumentParser(
        description="Academic Formula Inserter — Insert & format equations in Word/LaTeX"
    )
    sub = parser.add_subparsers(dest="command", help="Available commands")

    p_insert = sub.add_parser("insert-formula", help="Insert LaTeX formula into Word document")
    p_insert.add_argument("docx_path", help="Path to .docx file")
    p_insert.add_argument("--latex", required=True, help="LaTeX formula string")
    p_insert.add_argument("--display", action="store_true", help="Display equation (centered)")
    p_insert.add_argument("--style", help="Numbering style: chinese, parentheses, brackets, etc.")
    p_insert.add_argument("--number", type=int, help="Equation number")
    p_insert.add_argument("--position", type=int, help="Paragraph index to insert at")
    p_insert.add_argument("--label", help="Equation label for cross-reference")
    p_insert.add_argument("--interactive", action="store_true",
                          help="Interactive mode: prompt to choose numbering style")
    p_insert.add_argument("--mathtype", action="store_true",
                          help="MathType-compatible mode: MTDisplayEquation style + SEQ MTEqn")

    p_batch = sub.add_parser("batch-convert", help="Convert all $...$ LaTeX in document to OMML")
    p_batch.add_argument("docx_path", help="Path to .docx file")
    p_batch.add_argument("--backup", action="store_true", help="Create backup before conversion")
    p_batch.add_argument("--mode", choices=["quick", "safe", "full"], default="safe",
                         help="Conversion mode")

    p_number = sub.add_parser("number-formulas", help="Auto-number all equations")
    p_number.add_argument("docx_path", help="Path to .docx file")
    p_number.add_argument("--style", default="parentheses",
                          help="Numbering style: chinese, parentheses, brackets, section")
    p_number.add_argument("--interactive", action="store_true",
                          help="Interactive mode: prompt to choose numbering style")

    p_validate = sub.add_parser("validate-format", help="Validate formula format for journal")
    p_validate.add_argument("docx_path", help="Path to .docx file")
    p_validate.add_argument("--journal", required=True, help="Journal key (e.g. ieee, chinese-core)")

    p_export = sub.add_parser("export-latex", help="Export equations as LaTeX")
    p_export.add_argument("docx_path", help="Path to .docx file")
    p_export.add_argument("--output", help="Output .tex file path")

    sub.add_parser("list-styles", help="List all supported journal styles")

    args = parser.parse_args()

    if args.command == "insert-formula":
        cmd_insert_formula(args)
    elif args.command == "batch-convert":
        cmd_batch_convert(args)
    elif args.command == "number-formulas":
        cmd_number_formulas(args)
    elif args.command == "validate-format":
        cmd_validate_format(args)
    elif args.command == "export-latex":
        cmd_export_latex(args)
    elif args.command == "list-styles":
        cmd_list_styles(args)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
