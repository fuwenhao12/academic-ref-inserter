"""Create minimal test docx with one OMML formula, testing multiple approaches."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from docx.oxml import parse_xml
from lxml import etree

from latex_omml_converter import latex_to_omml

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ─── Approach 1: Manually crafted OMML matching working format ───
doc1 = Document()
doc1.add_heading('Approach 1: Manual OMML (matching working restored format)', level=2)
p = doc1.add_paragraph()
p.add_run('Taylor series: ')

manual_omml = (
    f'<m:oMath xmlns:m="{M_NS}">'
    f'<m:sSup>'
    f'<m:e><m:r><m:rPr><m:ital/></m:rPr><m:t>c</m:t></m:r></m:e>'
    f'<m:lim><m:r><m:t>2</m:t></m:r></m:lim>'
    f'</m:sSup>'
    f'</m:oMath>'
)
p._element.append(parse_xml(manual_omml))

# ─── Approach 2: Our converter output ───
doc2 = Document()
doc2.add_heading('Approach 2: Our latex_to_omml converter', level=2)
p2 = doc2.add_paragraph()
p2.add_run('Taylor series: ')
omml = latex_to_omml("c^2")
p2._element.append(parse_xml(omml))

# ─── Approach 3: With oMathPara wrapper ───
doc3 = Document()
doc3.add_heading('Approach 3: oMathPara wrapper', level=2)
p3 = doc3.add_paragraph()
p3.add_run('Taylor series: ')
omml3 = latex_to_omml("c^2")  # just <m:oMath>
elem3 = parse_xml(omml3)
omath_para3 = etree.SubElement(p3._element, f"{{{M_NS}}}oMathPara")
omath_para3.append(elem3)

# ─── Approach 4: Use `m:oMathPara` directly ───
doc4 = Document()
doc4.add_heading('Approach 4: Direct oMathPara', level=2)
p4 = doc4.add_paragraph()
p4.add_run('Taylor series: ')
omml4 = latex_to_omml("c^2", display=True)
p4._element.append(parse_xml(omml4))

# Save all
out_dir = Path(__file__).parent
doc1.save(str(out_dir / "test_approach1.docx"))
doc2.save(str(out_dir / "test_approach2.docx"))
doc3.save(str(out_dir / "test_approach3.docx"))
doc4.save(str(out_dir / "test_approach4.docx"))
print("All 4 test documents created!")

# Also print the OMML for comparison
print(f"\nApproach 1 OMML: {manual_omml}")
print(f"\nApproach 2 OMML: {omml}")
print(f"\nApproach 3 OMML (inner): {omml3}")
print(f"\nApproach 4 OMML: {omml4}")
