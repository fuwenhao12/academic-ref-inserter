from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = Document()

doc.add_heading('论文公式插入测试文档', level=1)

doc.add_paragraph(
    '本文档用于验证 academic-formula-inserter skill 的各项功能。'
    '包含 LaTeX 公式、待编号公式等测试内容。'
)

doc.add_heading('1. 基础公式', level=2)
doc.add_paragraph(
    '根据爱因斯坦的相对论，质能关系可以表示为 E=mc\u00b2，'
    '其中 E 是能量，m 是质量，c 是光速。'
)

doc.add_paragraph(
    '泰勒展开公式是数学分析中最重要的公式之一。'
)

doc.add_heading('2. 复杂公式', level=2)
doc.add_paragraph(
    '高斯积分公式是概率论和统计学中的基础公式。'
)

doc.add_paragraph(
    '神经网络中常用的激活函数为 sigmoid 函数，'
    '该函数将输入映射到 (0,1) 区间。'
)

doc.add_paragraph(
    '多尺度时序预测模型的损失函数由均方误差和正则化项组成。'
)

doc.add_heading('3. 待插入公式的占位段落', level=2)
docs_para = doc.add_paragraph()
docs_para.add_run('时序门控函数定义段落（公式将插入于此）').bold = True

var_para = doc.add_paragraph()
var_para.add_run('变量门控函数定义段落（公式将插入于此）').bold = True

fuse_para = doc.add_paragraph()
fuse_para.add_run('融合输出定义段落（公式将插入于此）').bold = True

doc.add_heading('4. 摘要', level=2)
doc.add_paragraph(
    '本文提出了一种面向多尺度时序预测的时域门控混合模型（HybridModel）。'
    '实验结果表明，所提方法在多个数据集上取得了最优性能。'
)

test_path = Path(__file__).parent / "test_output.docx"
doc.save(str(test_path))
print(f"Test document created: {test_path}")
