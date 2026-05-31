"""MathType integration for Word documents.

Creates MathType-compatible equations in Word .docx files using:
1. MTDisplayEquation paragraph style
2. OMML (Office Math Markup Language) for equation content
3. SEQ MTEqn field codes for equation numbering
4. MACROBUTTON MTPlaceRef for MathType reference placeholders

This makes equations double-click editable in MathType when installed.
"""
import io, zipfile, re, os, random
from pathlib import Path
from lxml import etree

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
V_NS = "urn:schemas-microsoft-com:vml"
O_NS = "urn:schemas-microsoft-com:office:office"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"

NSMAP = {
    'w': W_NS, 'm': M_NS, 'r': R_NS, 'mc': MC_NS,
    'w14': W14_NS, 'v': V_NS, 'o': O_NS, 'wp': WP_NS, 'a': A_NS,
}

MT_DISPLAY_EQUATION_STYLE_XML = """<w:style w:type="paragraph" w:customStyle="1" w:styleId="MTDisplayEquation"
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">
  <w:name w:val="MTDisplayEquation"/>
  <w:basedOn w:val="a1"/>
  <w:next w:val="a1"/>
  <w:pPr>
    <w:tabs>
      <w:tab w:val="center" w:pos="4820"/>
      <w:tab w:val="right" w:pos="9640"/>
    </w:tabs>
    <w:ind w:firstLine="420"/>
    <w:jc w:val="both"/>
  </w:pPr>
  <w:rPr>
    <w:lang w:eastAsia="zh-CN"/>
  </w:rPr>
</w:style>"""


def _random_hex(length=8):
    return ''.join(random.choices('0123456789ABCDEF', k=length))


def add_mt_display_equation_style(docx_path):
    """Add MTDisplayEquation style to the document's styles.xml."""
    with zipfile.ZipFile(docx_path, 'r') as zf:
        styles_xml = zf.read('word/styles.xml').decode('utf-8')

    if 'MTDisplayEquation' in styles_xml:
        return

    root = etree.fromstring(styles_xml.encode('utf-8'))
    style_elem = etree.fromstring(MT_DISPLAY_EQUATION_STYLE_XML.encode('utf-8'))
    root.append(style_elem)

    modified = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

    buffer = io.BytesIO()
    with zipfile.ZipFile(docx_path, 'r') as zf:
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as out:
            for item in zf.infolist():
                if item.filename == 'word/styles.xml':
                    out.writestr(item, modified)
                else:
                    out.writestr(item, zf.read(item.filename))

    with open(docx_path, 'wb') as f:
        f.write(buffer.getvalue())


def build_mt_equation_paragraph(omml_xml: str, equation_number: int = 1) -> str:
    """Build a complete MathType equation paragraph XML.

    Creates a paragraph with:
    - MTDisplayEquation style
    - OMML equation content
    - SEQ MTEqn field for numbering

    Matches the structure found in MathType-created documents.
    """
    para_id = _random_hex(8)
    text_id = _random_hex(8)
    rsid = _random_hex(8)

    omml_clean = omml_xml.strip()
    if omml_clean.startswith('<m:oMathPara'):
        omml_clean = omml_clean.replace(
            '<m:oMathPara', '<m:oMathPara'
        )

    return (
        f'<w:p w14:paraId="{para_id}" w14:textId="{text_id}"'
        f' w:rsidR="{rsid}" w:rsidRDefault="{rsid}"'
        f' xmlns:w="{W_NS}" xmlns:m="{M_NS}" xmlns:w14="{W14_NS}">'
        f'<w:pPr><w:pStyle w:val="MTDisplayEquation"/></w:pPr>'
        f'{omml_clean}'
        f'<w:fldSimple w:instr=" SEQ MTEqn \\\\c \\\\* Arabic \\\\* MERGEFORMAT ">'
        f'<w:r><w:rPr><w:noProof/></w:rPr>'
        f'<w:t>{equation_number}</w:t></w:r>'
        f'</w:fldSimple>'
        f'</w:p>'
    )


def build_mt_ole_paragraph(latex_text: str, equation_number: int = 1) -> str:
    """Build a MathType equation paragraph with MACROBUTTON MTPlaceRef.

    This creates a simpler placeholder that MathType can convert.
    When MathType is installed, it can convert OMML or plain text
    into full MathType equations.

    For true OLE embedding, the binary OLE file needs to exist.
    This creates a text-based placeholder that MathType can process.
    """
    para_id = _random_hex(8)
    text_id = _random_hex(8)
    rsid = _random_hex(8)

    return (
        f'<w:p w14:paraId="{para_id}" w14:textId="{text_id}"'
        f' w:rsidR="{rsid}" w:rsidRDefault="{rsid}"'
        f' xmlns:w="{W_NS}" xmlns:m="{M_NS}" xmlns:w14="{W14_NS}">'
        f'<w:pPr><w:pStyle w:val="MTDisplayEquation"/></w:pPr>'
        f'<w:r><w:rPr><w:lang w:eastAsia="zh-CN"/></w:rPr>'
        f'<w:fldChar w:fldCharType="begin"/></w:r>'
        f'<w:r><w:rPr><w:lang w:eastAsia="zh-CN"/></w:rPr>'
        f'<w:instrText xml:space="preserve"> MACROBUTTON MTPlaceRef \\\\* MERGEFORMAT </w:instrText></w:r>'
        f'<w:r><w:rPr><w:lang w:eastAsia="zh-CN"/></w:rPr>'
        f'<w:fldChar w:fldCharType="begin"/></w:r>'
        f'<w:r><w:rPr><w:lang w:eastAsia="zh-CN"/></w:rPr>'
        f'<w:instrText xml:space="preserve"> SEQ MTEqn \\\\h \\\\* MERGEFORMAT </w:instrText></w:r>'
        f'<w:r><w:rPr><w:lang w:eastAsia="zh-CN"/></w:rPr>'
        f'<w:fldChar w:fldCharType="end"/></w:r>'
        f'<w:r><w:rPr><w:lang w:eastAsia="zh-CN"/></w:rPr>'
        f'<w:instrText>(</w:instrText></w:r>'
        f'<w:r><w:rPr><w:lang w:eastAsia="zh-CN"/></w:rPr>'
        f'<w:instrText>)</w:instrText></w:r>'
        f'<w:r><w:rPr><w:lang w:eastAsia="zh-CN"/></w:rPr>'
        f'<w:fldChar w:fldCharType="end"/></w:r>'
        f'<w:fldSimple w:instr=" SEQ MTEqn \\\\c \\\\* Arabic \\\\* MERGEFORMAT ">'
        f'<w:r><w:rPr><w:noProof/></w:rPr>'
        f'<w:t>{equation_number}</w:t></w:r>'
        f'</w:fldSimple>'
        f'</w:p>'
    )


def build_mt_omml_equation(omml_xml: str, equation_number: int = 1, para_id: str = None) -> str:
    """Build a MathType paragraph using OMML (the modern approach).

    This is the recommended approach as it:
    - Produces native Word equations that render correctly
    - Allows double-click editing in Word (without MathType)
    - Opens in MathType when MathType is installed
    - Requires no binary OLE files
    """
    if para_id is None:
        para_id = _random_hex(8)
    text_id = _random_hex(8)
    rsid = _random_hex(8)

    return (
        f'<w:p w14:paraId="{para_id}" w14:textId="{text_id}"'
        f' w:rsidR="{rsid}" w:rsidRDefault="{rsid}"'
        f' xmlns:w="{W_NS}" xmlns:m="{M_NS}" xmlns:w14="{W14_NS}">'
        f'<w:pPr><w:pStyle w:val="MTDisplayEquation"/></w:pPr>'
        f'{omml_xml}'
        f'<w:fldSimple w:instr=" SEQ MTEqn \\\\c \\\\* Arabic \\\\* MERGEFORMAT ">'
        f'<w:r><w:rPr><w:noProof/></w:rPr>'
        f'<w:t>{equation_number}</w:t></w:r>'
        f'</w:fldSimple>'
        f'</w:p>'
    )


def inject_mt_paragraphs_via_zip(docx_path: str, paragraphs_xml: list):
    """Inject MathType-compatible paragraphs into a docx at the ZIP level.

    Args:
        docx_path: Path to the .docx file
        paragraphs_xml: List of paragraph XML strings (from build_mt_omml_equation)
    """
    with zipfile.ZipFile(docx_path, 'r') as zf:
        doc_xml = zf.read('word/document.xml')

    root = etree.fromstring(doc_xml)
    body = root.find(f'{{{W_NS}}}body')

    for para_xml in paragraphs_xml:
        para_elem = etree.fromstring(para_xml.encode('utf-8'))
        body.append(para_elem)

    modified = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

    buffer = io.BytesIO()
    with zipfile.ZipFile(docx_path, 'r') as zf:
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as out:
            for item in zf.infolist():
                if item.filename == 'word/document.xml':
                    out.writestr(item, modified)
                else:
                    out.writestr(item, zf.read(item.filename))

    with open(docx_path, 'wb') as f:
        f.write(buffer.getvalue())


def create_mathtype_document(output_path: str, formulas: list):
    """Create a complete Word document with MathType-compatible equations.

    Args:
        output_path: Path for the output .docx file
        formulas: List of (latex_str, is_display) tuples
    
    This is a convenience function that:
    1. Creates a base document using python-docx
    2. Adds MTDisplayEquation style to styles.xml
    3. Injects MathType-compatible paragraphs at ZIP level
    """
    from docx import Document
    from latex_omml_converter import latex_to_omml

    doc = Document()
    doc.add_heading('MathType 公式插入演示', level=1)
    doc.add_paragraph('以下公式使用 MathType 兼容格式，双击可在 MathType 中编辑。')

    temp_path = str(Path(output_path).parent / f"_temp_base_{_random_hex(4)}.docx")
    doc.save(temp_path)

    add_mt_display_equation_style(temp_path)

    paragraphs = []
    for i, (latex, is_display) in enumerate(formulas):
        omml = latex_to_omml(latex, display=False)
        para = build_mt_omml_equation(omml, i + 1)
        paragraphs.append(para)

    inject_mt_paragraphs_via_zip(temp_path, paragraphs)

    # Clean up python-docx corrupted elements
    with zipfile.ZipFile(temp_path, 'r') as zf:
        doc_xml = zf.read('word/document.xml')

    cleaned = re.sub(r'<m:sSupPr>.*?</m:sSupPr>', '', doc_xml.decode('utf-8'), flags=re.DOTALL)
    cleaned = re.sub(r'<m:sSubPr>.*?</m:sSubPr>', '', cleaned, flags=re.DOTALL)
    cleaned = re.sub(r'<m:ctrlPr>.*?</m:ctrlPr>', '', cleaned, flags=re.DOTALL)
    cleaned = re.sub(r'<m:sup/>', '', cleaned)
    cleaned = re.sub(r'<m:sub/>', '', cleaned)

    buffer = io.BytesIO()
    with zipfile.ZipFile(temp_path, 'r') as zf:
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as out:
            for item in zf.infolist():
                if item.filename == 'word/document.xml':
                    out.writestr(item, cleaned.encode('utf-8'))
                else:
                    out.writestr(item, zf.read(item.filename))

    with open(output_path, 'wb') as f:
        f.write(buffer.getvalue())

    Path(temp_path).unlink()
    return output_path


def analyze_mathtype_docx(docx_path: str):
    """Analyze a docx file for MathType-compatible formulas."""
    from latex_omml_converter import detect_equations_in_docx

    with zipfile.ZipFile(docx_path, 'r') as zf:
        doc_xml = zf.read('word/document.xml').decode('utf-8')
        styles_xml = zf.read('word/styles.xml').decode('utf-8')

    results = {
        'has_mt_style': 'MTDisplayEquation' in styles_xml,
        'has_seq_mteqn': 'SEQ MTEqn' in doc_xml,
        'has_mt_place_ref': 'MTPlaceRef' in doc_xml,
        'total_omml': len(re.findall(r'<m:oMath[ >]', doc_xml)),
        'total_ole': len(re.findall(r'Equation\.DSMT4', doc_xml)),
        'total_fld_simple': len(re.findall(r'<w:fldSimple[ >]', doc_xml)),
    }

    return results


if __name__ == '__main__':
    import sys
    formulas = [
        ("E=mc^2", False),
        ("\\sum_{i=1}^{n} x_i", True),
        ("\\frac{1}{1+e^{-x}}", True),
    ]

    output = r'c:\Users\HP\Desktop\汇总\submission_package_computer_engineering\submission_package\.trae\skills\academic-formula-inserter\scripts\mt_test.docx'
    create_mathtype_document(output, formulas)
    print(f"Created: {output}")
    result = analyze_mathtype_docx(output)
    print(f"Analysis: {result}")
