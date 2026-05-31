"""创建带真实 OMML 公式的演示文档"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
sys.path.insert(0, str(Path(__file__).parent / "utils"))

from docx import Document
from docx.oxml import parse_xml
from docx.shared import Pt
from lxml import etree

from latex_omml_converter import latex_to_omml
from formula_numbering import _add_number_to_paragraph

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

doc = Document()

doc.add_heading('学术公式插入演示', level=1)
doc.add_paragraph('以下公式使用 academic-formula-inserter 的 LaTeX→OMML 转换引擎插入，可在 Word 中直接双击编辑。')

# ─── 公式1: E=mc^2（行内，圆括号编号） ───
doc.add_heading('1. 质能方程', level=2)
p1 = doc.add_paragraph()
p1.add_run('爱因斯坦质能关系：')
omml1 = latex_to_omml("E=mc^2")
elem1 = parse_xml(omml1)
p1._element.append(elem1)
_add_number_to_paragraph(p1, "(1)")

# ─── 公式2: 求和公式（显示，中文编号） ───
doc.add_heading('2. 求和公式', level=2)
p2 = doc.add_paragraph()
p2.add_run('离散求和定义：')
omml2 = latex_to_omml("\\sum_{i=1}^{n} x_i")
elem2 = parse_xml(omml2)
omath_para2 = etree.SubElement(p2._element, f"{{{M_NS}}}oMathPara")
omath_para2.append(elem2)
_add_number_to_paragraph(p2, "式(2)")

# ─── 公式3: 分数+指数（显示，方括号编号） ───
doc.add_heading('3. Sigmoid 激活函数', level=2)
p3 = doc.add_paragraph()
p3.add_run('Sigmoid 函数定义：')
omml3 = latex_to_omml("\\frac{1}{1+e^{-x}}")
elem3 = parse_xml(omml3)
omath_para3 = etree.SubElement(p3._element, f"{{{M_NS}}}oMathPara")
omath_para3.append(elem3)
_add_number_to_paragraph(p3, "[3]")

# ─── 公式4: 门控函数（显示，中文编号） ───
doc.add_heading('4. 时序门控函数', level=2)
p4 = doc.add_paragraph()
p4.add_run('时域门控混合模型的核心公式：')
omml4 = latex_to_omml("g_t = \\sigma(W_t x_t + b_t)")
elem4 = parse_xml(omml4)
omath_para4 = etree.SubElement(p4._element, f"{{{M_NS}}}oMathPara")
omath_para4.append(elem4)
_add_number_to_paragraph(p4, "式(4)")

# ─── 公式5: 复杂公式（显示，圆括号编号） ───
doc.add_heading('5. 损失函数', level=2)
p5 = doc.add_paragraph()
p5.add_run('多尺度时序预测模型损失函数：')
omml5 = latex_to_omml("L = \\frac{1}{N}\\sum_{i=1}^{N}(y_i - \\hat{y}_i)^2 + \\lambda\\|\\theta\\|_2^2")
elem5 = parse_xml(omml5)
omath_para5 = etree.SubElement(p5._element, f"{{{M_NS}}}oMathPara")
omath_para5.append(elem5)
_add_number_to_paragraph(p5, "(5)")

demo_path = Path(__file__).parent / "formula_demo_v3.docx"
doc.save(str(demo_path))
print(f"✅ 演示文档已创建: {demo_path}")
print(f"   共插入 5 个 OMML 公式（3 种编号风格）")
