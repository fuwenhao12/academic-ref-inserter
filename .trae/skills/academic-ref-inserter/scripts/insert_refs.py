#!/usr/bin/env python3
"""
Academic Reference Inserter
Main entry point for inserting, formatting, and cross-referencing
academic references in Word documents.

Supports human-readable (default) and JSON (--json) output modes
for compatibility with AI coding assistants.

Usage:
    python insert_refs.py --version
    python insert_refs.py analyze --input paper.docx
    python insert_refs.py reformat --input paper.docx --format gbt7714
    python insert_refs.py reorder --input paper.docx
    python insert_refs.py hyperlink --input paper.docx
    python insert_refs.py validate --input paper.docx --format gbt7714
    python insert_refs.py fix --input paper.docx --format gbt7714 --json
"""

import os
import sys
import json
import re
import argparse

from formats.gbt7714 import GBT7714Format
from formats.ieee import IEEEFormat
from formats.apa7 import APA7Format
from formats.chicago import ChicagoFormat
from formats.mla import MLAFormat
from formats.harvard import HarvardFormat
from utils import docx_utils
from utils.parser import parse_ref, detect_type
from utils.validator import validate_all
from utils.doi_lookup import lookup_doi, lookup_by_title, format_from_doi
from utils.bibtex import parse_bibtex, export_bibtex

__version__ = "1.0.0"


FORMATS = {
    'gbt7714': GBT7714Format(),
    'ieee': IEEEFormat(),
    'apa7': APA7Format(),
    'chicago': ChicagoFormat(),
    'mla': MLAFormat(),
    'harvard': HarvardFormat(),
}


def _get_format(name):
    fmt = FORMATS.get(name.lower())
    if not fmt:
        raise ValueError(f"Unsupported format: {name}. Available: {', '.join(FORMATS.keys())}")
    return fmt


def _emit(result, args):
    """Output result: human-readable if no --json, JSON otherwise."""
    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    elif 'message' in result:
        print(result['message'])


def _save_doc(doc, args):
    """Save document to --output if specified, otherwise overwrite --input."""
    target = args.output if args.output else args.input
    doc.save(target)
    return target


def _analysis_to_json(args, doc, ref_title_idx, ref_entries, citation_order):
    """Build structured analysis result."""
    ref_nums = {num for _, _, num in ref_entries}
    cited_nums = set(citation_order)
    max_ref = max(ref_nums) if ref_nums else (max(citation_order) if citation_order else 0)
    expected = list(range(1, max_ref + 1)) if max_ref > 0 else []
    orphan = sorted(ref_nums - cited_nums)
    missing = sorted(cited_nums - ref_nums)

    ref_list = []
    for idx, para, num in sorted(ref_entries, key=lambda x: x[2]):
        text = para.text.strip()
        ctx = docx_utils.extract_citation_context(doc, num)
        ref_list.append({
            'number': num,
            'type': detect_type(text),
            'text': text,
            'context_snippet': ctx,
        })

    return {
        'command': 'analyze',
        'input_file': args.input,
        'ref_section_paragraph': ref_title_idx,
        'total_refs': len(ref_entries),
        'citation_order': citation_order,
        'sequential_ok': citation_order == expected,
        'expected_sequential': expected,
        'references': ref_list,
        'orphan_refs': orphan,
        'missing_refs': missing,
        'has_orphan': len(orphan) > 0,
        'has_missing': len(missing) > 0,
        'file_ok': len(orphan) == 0 and len(missing) == 0,
    }


def cmd_analyze(args):
    """Analyze document: extract citations, references, and detect issues."""
    doc = docx_utils.open_docx(args.input)
    ref_title_idx, ref_entries = docx_utils.extract_reference_section(doc)
    citation_order = docx_utils.extract_citations_from_text(doc)

    result = _analysis_to_json(args, doc, ref_title_idx, ref_entries, citation_order)
    result['message'] = (
        f"Document: {args.input}\n"
        f"Reference section at P{ref_title_idx}\n"
        f"Total refs: {len(ref_entries)}\n"
        f"Citation order: {' -> '.join(str(n) for n in citation_order)}\n"
        f"Sequential: {'OK' if result['sequential_ok'] else 'NEEDS REORDER'}\n"
        f"Orphan refs: {result['orphan_refs'] or 'none'}\n"
        f"Missing refs: {result['missing_refs'] or 'none'}"
    )
    _emit(result, args)

    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)


def cmd_reformat(args):
    """Reformat all reference entries to the target citation style."""
    fmt = _get_format(args.format)
    doc = docx_utils.open_docx(args.input)
    backup_path = docx_utils.backup_docx(args.input)

    ref_title_idx, ref_entries = docx_utils.extract_reference_section(doc)
    changes = []

    for idx, para, num in ref_entries:
        raw_text = para.text.strip()
        parsed = parse_ref(raw_text)
        new_ref = fmt.format(parsed)
        new_ref = f"[{num}] {new_ref}"

        if new_ref != raw_text:
            docx_utils.update_reference_text(para, new_ref)
            changes.append({'number': num, 'old': raw_text[:80], 'new': new_ref[:80]})

    _save_doc(doc, args)
    result = {
        'command': 'reformat',
        'input_file': args.input,
        'backup_file': backup_path,
        'format': args.format,
        'total_entries': len(ref_entries),
        'changes': len(changes),
        'messages': [f"[{c['number']}] reformatted" for c in changes],
    }
    result['message'] = f"Reformatted {len(changes)}/{len(ref_entries)} entries to {fmt.name}"
    _emit(result, args)


def cmd_reorder(args):
    """Reorder references to match first citation order in text."""
    doc = docx_utils.open_docx(args.input)
    backup_path = docx_utils.backup_docx(args.input)

    citation_order = docx_utils.extract_citations_from_text(doc)
    ref_title_idx, ref_entries = docx_utils.extract_reference_section(doc)

    if not ref_entries:
        result = {'command': 'reorder', 'error': 'No references found'}
        _emit(result, args)
        return

    old_to_new = {}
    for new_num, old_num in enumerate(citation_order, 1):
        old_to_new[old_num] = new_num

    # Update bib entries
    bib_changes = 0
    for idx, para, old_num in ref_entries:
        new_num = old_to_new.get(old_num, old_num)
        raw_text = para.text.strip()
        new_text = re.sub(r'^\[\d+\]', f'[{new_num}]', raw_text)
        if new_text != raw_text:
            docx_utils.update_reference_text(para, new_text)
            bib_changes += 1

    # Update in-text citations using two-phase approach to avoid cascading
    text_changes = 0
    for para in doc.paragraphs:
        for old_num, new_num in old_to_new.items():
            if old_num == new_num:
                continue
            for run in para.runs:
                old_str = f'[{old_num}]'
                if old_str in run.text:
                    # Phase 1: replace [old_num] with a safe temporary marker
                    marker = f'__REORDER_{new_num}__'
                    run.text = run.text.replace(old_str, marker)
                    text_changes += 1

    # Phase 2: replace all temporary markers with final format
    for para in doc.paragraphs:
        for new_num in set(old_to_new.values()):
            marker = f'__REORDER_{new_num}__'
            final = f'[{new_num}]'
            for run in para.runs:
                if marker in run.text:
                    run.text = run.text.replace(marker, final)

    _save_doc(doc, args)

    mapping = {str(k): str(v) for k, v in old_to_new.items() if k != v}
    result = {
        'command': 'reorder',
        'input_file': args.input,
        'backup_file': backup_path,
        'bib_entries_updated': bib_changes,
        'in_text_updated': text_changes,
        'renumbering_mapping': mapping,
    }
    result['message'] = f"Updated {bib_changes} bib entries, {text_changes} in-text citations"
    _emit(result, args)


def cmd_hyperlink(args):
    """Add bookmarks and hyperlinks from citations to bibliography entries."""
    doc = docx_utils.open_docx(args.input)
    backup_path = docx_utils.backup_docx(args.input)

    ref_title_idx, ref_entries = docx_utils.extract_reference_section(doc)

    bookmarks = 0
    for idx, para, num in ref_entries:
        if docx_utils.add_bookmark_to_paragraph(para, f"Ref_{num}"):
            bookmarks += 1

    hyperlinks = 0
    dedup_total = 0
    for i, para in enumerate(doc.paragraphs):
        if ref_title_idx is not None and i >= ref_title_idx:
            break
        dedup_total += docx_utils.dedup_adjacent_citations(para)
        refs = re.findall(r'\[(\d+)\]', para.text)
        if not refs:
            continue
        for ref_num in sorted(set(int(r) for r in refs), reverse=True):
            if docx_utils.replace_citation_with_hyperlink(para, ref_num):
                hyperlinks += 1

    _save_doc(doc, args)
    result = {
        'command': 'hyperlink',
        'input_file': args.input,
        'backup_file': backup_path,
        'bookmarks_added': bookmarks,
        'hyperlinks_created': hyperlinks,
        'total_refs': len(ref_entries),
    }
    result['message'] = f"Bookmarks: {bookmarks}, Hyperlinks: {hyperlinks}, Dedup: {dedup_total}"
    _emit(result, args)


def cmd_validate(args):
    """Validate reference consistency and format."""
    fmt = _get_format(args.format) if args.format else None
    doc = docx_utils.open_docx(args.input)

    ref_title_idx, ref_entries = docx_utils.extract_reference_section(doc)
    citation_order = docx_utils.extract_citations_from_text(doc)

    ref_texts = [para.text.strip() for _, para, _ in ref_entries]
    ref_nums = {num for _, _, num in ref_entries}
    cited_nums = set(citation_order)
    years = []

    for rt in ref_texts:
        parsed = parse_ref(rt)
        years.append(parsed.get('year', ''))

    style_name = fmt.name if fmt else 'gbt7714'
    report = validate_all(ref_texts, cited_nums, ref_nums, years, style=style_name)

    issues_detail = []
    for issue in report.get('issues', []):
        issues_detail.append({'severity': 'error', 'text': issue})
    for warning in report.get('warnings', []):
        issues_detail.append({'severity': 'warning', 'text': warning})

    result = {
        'command': 'validate',
        'input_file': args.input,
        'format': style_name,
        'total_refs': report['total_refs'],
        'passed': report['passed'],
        'issues': issues_detail,
        'recency': None,
    }

    if report.get('recency'):
        result['recency'] = {
            'recent_count': report['recency']['recent_count'],
            'total': report['recency']['total'],
            'ratio': round(report['recency']['ratio'], 2),
            'threshold_ok': report['recency']['ok'],
        }

    summary = f"Validation ({style_name}): {'PASSED' if report['passed'] else 'FAILED'}\n"
    for issue in issues_detail:
        summary += f"  [{issue['severity'].upper()}] {issue['text']}\n"
    result['message'] = summary.strip()
    _emit(result, args)

    if not report['passed'] and not args.json:
        sys.exit(1)
    elif not report['passed'] and args.json:
        result['exit_code'] = 1


def cmd_fix(args):
    """Full pipeline: analyze + reformat + reorder + hyperlink + validate."""
    fmt = _get_format(args.format)
    backup_path = docx_utils.backup_docx(args.input)

    steps = []

    # Step 1: Analyze
    doc = docx_utils.open_docx(args.input)
    ref_title_idx, ref_entries = docx_utils.extract_reference_section(doc)
    citation_order = docx_utils.extract_citations_from_text(doc)
    analysis = _analysis_to_json(args, doc, ref_title_idx, ref_entries, citation_order)
    steps.append({'step': 'analyze', 'total_refs': analysis['total_refs']})

    # Step 2: Reformat
    for idx, para, num in ref_entries:
        raw_text = para.text.strip()
        parsed = parse_ref(raw_text)
        new_ref = fmt.format(parsed)
        new_ref = f"[{num}] {new_ref}"
        if new_ref != raw_text:
            docx_utils.update_reference_text(para, new_ref)
    _save_doc(doc, args)
    steps.append({'step': 'reformat', 'format': args.format})

    # Step 3: Reorder (if needed)
    if not analysis['sequential_ok']:
        old_to_new = {}
        for new_num, old_num in enumerate(citation_order, 1):
            old_to_new[old_num] = new_num
        doc = docx_utils.open_docx(args.input)
        for idx, para, old_num in ref_entries:
            new_num = old_to_new.get(old_num, old_num)
            raw_text = para.text.strip()
            new_text = re.sub(r'^\[\d+\]', f'[{new_num}]', raw_text)
            if new_text != raw_text:
                docx_utils.update_reference_text(para, new_text)
        for para in doc.paragraphs:
            for old_num, new_num in old_to_new.items():
                if old_num == new_num:
                    continue
                if f'[{old_num}]' in para.text:
                    for run in para.runs:
                        if f'[{old_num}]' in run.text:
                            run.text = run.text.replace(f'[{old_num}]', f'[{new_num}]')
        _save_doc(doc, args)
        steps.append({'step': 'reorder', 'mapping': {str(k): str(v) for k, v in old_to_new.items() if k != v}})

    # Step 4: Hyperlinks
    doc = docx_utils.open_docx(args.input)
    ref_title_idx, ref_entries = docx_utils.extract_reference_section(doc)
    for idx, para, num in ref_entries:
        docx_utils.add_bookmark_to_paragraph(para, f"Ref_{num}")
    hyperlinks = 0
    for para in doc.paragraphs:
        refs = re.findall(r'\[(\d+)\]', para.text)
        if not refs:
            continue
        for ref_num in sorted(set(int(r) for r in refs), reverse=True):
            if docx_utils.replace_citation_with_hyperlink(para, ref_num):
                hyperlinks += 1
    _save_doc(doc, args)
    steps.append({'step': 'hyperlink', 'hyperlinks_created': hyperlinks})

    # Step 5: Validate
    ref_texts = [para.text.strip() for _, para, _ in ref_entries]
    ref_nums = {num for _, _, num in ref_entries}
    years = []
    for rt in ref_texts:
        parsed = parse_ref(rt)
        years.append(parsed.get('year', ''))
    report = validate_all(ref_texts, ref_nums, set(citation_order), years, style=fmt.name)
    steps.append({'step': 'validate', 'passed': report['passed']})

    result = {
        'command': 'fix',
        'input_file': args.input,
        'backup_file': backup_path,
        'format': args.format,
        'format_name': fmt.name,
        'steps': steps,
        'passed': report['passed'],
        'validation': {
            'total_refs': report['total_refs'],
            'issues': [i for i in report.get('issues', [])],
            'warnings': [w for w in report.get('warnings', [])],
        },
    }

    summary_lines = [
        f"Academic Ref Inserter - Full Pipeline ({fmt.name})",
        f"  [1/5] Analyze: {analysis['total_refs']} refs found",
        f"  [2/5] Reformat: {args.format}",
        f"  [3/5] Reorder: {'done' if not analysis['sequential_ok'] else 'skipped (already sequential)'}",
        f"  [4/5] Hyperlinks: {hyperlinks} created",
        f"  [5/5] Validate: {'PASSED' if report['passed'] else 'FAILED'}",
        f"  Backup: {backup_path}",
    ]
    result['message'] = '\n'.join(summary_lines)
    _emit(result, args)

    if not report['passed'] and not args.json:
        sys.exit(1)


def cmd_check_refs(args):
    """Check all references: auto-detect reference list, verify citations, report status.

    Provides a comprehensive per-reference citation check, showing
    which references are cited in the text and which are not.
    """
    doc = docx_utils.open_docx(args.input)

    # Step 1: Robustly find reference section
    ref_title_idx = docx_utils.find_reference_boundary_robust(doc)

    if ref_title_idx is None:
        # Fallback: try the standard method
        ref_title_idx, _ = docx_utils.extract_reference_section(doc)

    if ref_title_idx is None:
        err = {'error': 'Could not find reference section. Ensure the document has a "参考文献" or "References" header.'}
        print(json.dumps(err) if args.json else err['error'])
        sys.exit(1)

    # Step 2: Extract reference entries
    _, ref_entries = docx_utils.extract_reference_section(doc)

    if not ref_entries:
        err = {'error': 'No reference entries found. Ensure references are formatted as [1], [2], etc.'}
        print(json.dumps(err) if args.json else err['error'])
        sys.exit(1)

    # Step 3: Extract all in-text citations
    citation_order = docx_utils.extract_citations_from_text(doc)

    ref_nums = {num for _, _, num in ref_entries}
    cited_nums = set(citation_order)

    # Step 4: For each reference, find where it's cited
    ref_details = []
    for idx, para, num in sorted(ref_entries, key=lambda x: x[2]):
        text = para.text.strip()
        parsed = parse_ref(text)
        occurrences = docx_utils.find_citation_occurrences(doc, num)
        is_cited = num in cited_nums
        ref_details.append({
            'number': num,
            'type': parsed.get('type', 'unknown'),
            'text': text[:120] + ('...' if len(text) > 120 else ''),
            'cited': is_cited,
            'occurrence_count': len(occurrences),
            'occurrences': [{'para': o['para_index'], 'snippet': o['snippet']} for o in occurrences],
        })

    # Step 5: Find cited-in-text but not-in-bibliography
    all_cited_nums = set()
    for para in doc.paragraphs:
        for match in re.finditer(r'\[([0-9,\-\s]+)\]', para.text):
            content = match.group(1)
            for num in docx_utils.extract_citation_content(content):
                all_cited_nums.add(num)

    missing_nums = sorted(all_cited_nums - ref_nums)

    # Step 6: Build result
    cited_count = sum(1 for r in ref_details if r['cited'])
    uncited_count = sum(1 for r in ref_details if not r['cited'])

    result = {
        'command': 'check-refs',
        'input_file': args.input,
        'ref_section_at': f'P{ref_title_idx}',
        'total_refs': len(ref_entries),
        'cited_count': cited_count,
        'uncited_count': uncited_count,
        'missing_citations': missing_nums,
        'all_cited': uncited_count == 0 and len(missing_nums) == 0,
        'ref_details': ref_details,
    }

    # Human-readable output
    if not args.json:
        lines = [
            '=' * 60,
            '  Reference Citation Check Report',
            '=' * 60,
            f'  Reference section: P{ref_title_idx}',
            f'  Total refs: {len(ref_entries)}',
            f'  Cited: {cited_count}',
            f'  Uncited: {uncited_count}',
            f'  Missing from bib: {len(missing_nums)}',
            '',
        ]

        if missing_nums:
            lines.append(f'  ⚠  In-text citations with NO matching reference:')
            for n in missing_nums:
                lines.append(f'       [{n}]  ← CITED in text but NOT in reference list')
            lines.append('')

        lines.append(f'  {"=" * 56}')
        lines.append(f'  {"Ref":>4}  {"Cited":<6}  {"Type":<12}  {"Content":<60}')
        lines.append(f'  {"-" * 56}')

        for rd in ref_details:
            status = '[CITED]' if rd['cited'] else '[NOT CITED]'
            occ = f' (x{rd["occurrence_count"]})' if rd['occurrence_count'] > 1 else ''
            lines.append(f'  [{rd["number"]:>2}]  {status:<10}{occ:<8}  {rd["type"]:<12}  {rd["text"][:60]}')
            if rd['cited'] and rd['occurrences']:
                locations = [f"P{o['para']+1}" for o in rd['occurrences'][:3]]
                extra = f"  ... +{len(rd['occurrences'])-3} more" if len(rd['occurrences']) > 3 else ''
                lines.append(f'        at: {", ".join(locations)}{extra}')

        lines.append(f'  {"=" * 56}')

        if uncited_count > 0:
            lines.append('')
            lines.append(f'  [!] Uncited references (in bibliography but NOT cited in text):')
            for rd in ref_details:
                if not rd['cited']:
                    lines.append(f'       [{rd["number"]}]  {rd["text"][:70]}')
            lines.append('')
            lines.append('  Recommendation: Either add in-text citations or remove from bibliography.')

        if missing_nums:
            lines.append('')
            lines.append(f'  [!] Missing references (cited in text but NOT in bibliography):')
            for n in missing_nums:
                lines.append(f'       [{n}]')
            lines.append('')
            lines.append('  Recommendation: Add the missing reference entries to the bibliography.')

        status_icon = '[OK]' if result['all_cited'] else '[WARN]'
        lines.append('')
        lines.append(f'  Summary: {status_icon} All references checked.')
        if result['all_cited']:
            lines.append('  All references are properly cited in the text.')
        else:
            lines.append(f'  {uncited_count} uncited, {len(missing_nums)} missing - see details above.')
        lines.append('')

        result['message'] = '\n'.join(lines)
    else:
        result['message'] = f'Checked {len(ref_entries)} refs: {cited_count} cited, {uncited_count} uncited'

    # Export report to file if --report is given
    report_path = getattr(args, 'report', None)
    if report_path and lines:
        report_content = '\n'.join(lines)
        try:
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
            result['report_saved'] = report_path
            result['message'] += f'\n  Report saved to: {report_path}'
        except Exception as e:
            result['report_error'] = str(e)

    _emit(result, args)


def _extract_paper_topic(doc) -> str:
    """Extract the paper's title/topic from the first few paragraphs."""
    topic_parts = []
    for i, para in enumerate(doc.paragraphs):
        if i > 8:
            break
        text = para.text.strip()
        if not text:
            continue
        # Skip obvious non-title paragraphs
        if any(text.startswith(h) for h in ('Abstract', '摘要', 'Keywords', '关键词',
                                             '1 ', '1.', 'I.', '第一章')):
            continue
        # Prefer short, substantive text (likely a title)
        if 10 < len(text) < 120:
            topic_parts.append(text)
            # A real title rarely spans more than 2 paragraphs
            if len(topic_parts) >= 2:
                break
    topic = ' '.join(topic_parts)
    # Fallback: use first substantive paragraph
    if not topic:
        for para in doc.paragraphs:
            text = para.text.strip()
            if len(text) > 20:
                topic = text[:100]
                break
    return topic


def _search_and_pick(args, query: str, max_results: int) -> list:
    """Search CrossRef for refs matching query, let user pick, return formatted refs.

    In non-interactive mode (--json), returns all results.
    In interactive mode, prompts user to select which refs to include.
    Returns list of dicts: {doi, title, authors, year, formatted}.
    """
    results = lookup_by_title(query, max_results=max_results)
    if not results:
        return []

    fmt = _get_format(args.format)
    formatted_results = []
    for r in results:
        formatted_results.append({
            'doi': r.get('doi', ''),
            'title': r.get('title', '')[:100],
            'authors': ', '.join(r.get('authors', [])[:3]),
            'year': r.get('year', ''),
            'formatted': fmt.format(r),
        })

    # Deduplicate by DOI
    seen_dois = set()
    unique_results = []
    for r in formatted_results:
        if r['doi'] and r['doi'] not in seen_dois:
            seen_dois.add(r['doi'])
            unique_results.append(r)

    return unique_results


def _insert_new_refs(doc, new_entries: list) -> int:
    """Insert new formatted reference entries into the reference section.

    Appends them after the last existing reference entry.
    Returns the number of entries inserted.
    """
    from lxml import etree
    from docx.oxml.ns import qn

    ref_title_idx, ref_entries = docx_utils.extract_reference_section(doc)

    if ref_title_idx is None:
        return 0

    if not ref_entries:
        insert_idx = ref_title_idx + 1
        next_num = 1
    else:
        insert_idx = ref_entries[-1][0] + 1
        next_num = ref_entries[-1][2] + 1

    body = doc.element.body
    inserted = 0
    for i, entry in enumerate(new_entries):
        num = next_num + i
        full_text = f'[{num}] {entry["formatted"]}'

        new_p = etree.SubElement(body, qn('w:p'))
        new_r = etree.SubElement(new_p, qn('w:r'))
        new_t = etree.SubElement(new_r, qn('w:t'))
        new_t.text = full_text

        # Move new_p to correct position
        paras = list(doc.paragraphs)
        if insert_idx + inserted < len(paras):
            ref_para = paras[insert_idx + inserted]._element
            ref_para.addnext(new_p)
        else:
            body.append(new_p)

        inserted += 1

    return inserted


def cmd_auto_find(args):
    """Auto-search for relevant academic references and insert them into the document.

    Extracts the paper's title/topic, searches CrossRef for related works,
    presents candidates for user selection, and inserts chosen references
    into the document's reference section with proper formatting.
    """
    doc = docx_utils.open_docx(args.input)

    # Step 1: Extract paper topic
    topic = args.query if args.query else _extract_paper_topic(doc)
    if not topic:
        err = {'error': 'Could not extract paper topic. Provide --query or ensure the document has a title.'}
        print(json.dumps(err) if args.json else err['error'])
        sys.exit(1)

    # Step 2: Search for references
    max_results = args.max_results if args.max_results else 10
    candidates = _search_and_pick(args, topic, max_results)

    if not candidates:
        err = {'error': f'No references found for query: {topic[:60]}'}
        print(json.dumps(err) if args.json else err['error'])
        sys.exit(1)

    # Step 3: Present results
    if not args.json:
        lines = [
            f'Found {len(candidates)} references for: {topic[:80]}',
            '',
        ]
        for i, c in enumerate(candidates):
            lines.append(f'  [{i+1}] {c["formatted"][:90]}')
            lines.append(f'       DOI: {c["doi"][:60]}')
        lines.append('')
        lines.append('Enter numbers to select (e.g. "1,3-5"), or "all" for all, or "none" to skip:')
        print('\n'.join(lines))

        if not getattr(args, 'no_interactive', False):
            try:
                choice = input('> ').strip().lower()
            except (EOFError, KeyboardInterrupt):
                choice = 'none'
                print()

            if choice in ('none', 'n', ''):
                selected = []
                print('No references selected.')
            elif choice == 'all':
                selected = candidates
            else:
                selected = []
                for part in choice.split(','):
                    part = part.strip()
                    if '-' in part:
                        a, _, b = part.partition('-')
                        try:
                            for n in range(int(a.strip()), int(b.strip()) + 1):
                                if 1 <= n <= len(candidates):
                                    selected.append(candidates[n - 1])
                        except ValueError:
                            pass
                    else:
                        try:
                            n = int(part)
                            if 1 <= n <= len(candidates):
                                selected.append(candidates[n - 1])
                        except ValueError:
                            pass
        else:
            selected = candidates
    else:
        # Non-interactive: all candidates
        selected = candidates

    # Step 4: Insert into document (only if selected)
    if not selected:
        result = {
            'command': 'auto-find',
            'input_file': args.input,
            'query': topic[:80],
            'candidates_found': len(candidates),
            'selected': 0,
            'inserted': 0,
        }
        result['message'] = f'Found {len(candidates)} refs, selected 0.'
        _emit(result, args)
        return

    # Filter out refs already in the document
    ref_title_idx, ref_entries = docx_utils.extract_reference_section(doc)
    existing_dois = set()
    existing_titles = set()
    if ref_entries:
        for _, para, _ in ref_entries:
            text = para.text.lower().strip()
            existing_dois.update(re.findall(r'10\.\d{4,}/[^\s\]]+', text))
            # Extract short title fragments for dedup
            words = set(text.split()[:5])
            existing_titles.update(words)

    new_to_add = []
    for c in selected:
        doi = c.get('doi', '')
        if doi in existing_dois:
            continue
        # Check title overlap
        title_words = set(c.get('title', '').lower().split()[:5])
        if title_words and title_words.issubset(existing_titles):
            continue
        new_to_add.append(c)

    if not new_to_add:
        # Save anyway
        _save_doc(doc, args)
        result = {
            'command': 'auto-find',
            'input_file': args.input,
            'query': topic[:80],
            'candidates_found': len(candidates),
            'selected': len(selected),
            'inserted': 0,
            'message': f'Found {len(candidates)} refs, selected {len(selected)}, but all already exist in bibliography.',
        }
        _emit(result, args)
        return

    inserted_count = _insert_new_refs(doc, new_to_add)

    # Auto-reorder after insertion to fix all citation numbers
    if inserted_count > 0:
        print(f'\n  Auto-reordering citations after inserting {inserted_count} new references...')
        reorder_result = cmd_reorder(args)
        print(f'  Reorder complete.')
    else:
        saved_path = _save_doc(doc, args)

    if inserted_count > 0:
        saved_path = _save_doc(doc, args)

    result = {
        'command': 'auto-find',
        'input_file': args.input,
        'query': topic[:80],
        'candidates_found': len(candidates),
        'selected': len(selected),
        'existing_skipped': len(selected) - len(new_to_add),
        'inserted': inserted_count,
        'auto_reordered': inserted_count > 0,
        'saved_to': saved_path,
    }
    reorder_msg = '\n  Auto-reordered citations to fix numbering.' if inserted_count > 0 else ''
    result['message'] = (
        f'Auto-find results for: {topic[:80]}\n'
        f'  Candidates found: {len(candidates)}\n'
        f'  Selected: {len(selected)}\n'
        f'  Already in bib (skipped): {len(selected) - len(new_to_add)}\n'
        f'  Inserted: {inserted_count}\n'
        f'  Saved to: {saved_path}{reorder_msg}'
    )
    _emit(result, args)


def cmd_generate(args):
    """Generate a formatted reference from DOI, title search, or BibTeX input."""
    if args.doi:
        doi = args.doi.strip()
        formatted = format_from_doi(doi, args.format)
        if formatted:
            result = {'command': 'generate', 'source': 'doi', 'doi': doi, 'formatted': formatted}
            result['message'] = formatted
            _emit(result, args)
        else:
            err = {'error': f'Failed to look up DOI: {doi}'}
            print(json.dumps(err) if args.json else err['error'])
            sys.exit(1)
    elif args.search:
        results = lookup_by_title(args.search, max_results=5)
        if results:
            fmt = _get_format(args.format)
            formatted_list = []
            for r in results:
                formatted_list.append({
                    'doi': r.get('doi', ''),
                    'title': r.get('title', '')[:80],
                    'authors': ', '.join(r.get('authors', [])[:3]),
                    'year': r.get('year', ''),
                    'formatted': fmt.format(r),
                })
            result = {
                'command': 'generate',
                'source': 'search',
                'query': args.search,
                'results': formatted_list,
            }
            msg = '\n'.join(f"  [{i+1}] {r['formatted']}" for i, r in enumerate(formatted_list))
            result['message'] = f"Search: {args.search}\n{msg}"
            _emit(result, args)
        else:
            err = {'error': f'No results for: {args.search}'}
            print(json.dumps(err) if args.json else err['error'])
            sys.exit(1)
    else:
        err = {'error': 'Specify --doi or --search'}
        print(json.dumps(err) if args.json else err['error'])
        sys.exit(1)


def cmd_bib(args):
    """Import from or export to BibTeX format."""
    if args.import_file:
        with open(args.import_file, 'r', encoding='utf-8') as f:
            bib_content = f.read()
        entries = parse_bibtex(bib_content)
        fmt = _get_format(args.format)
        formatted = []
        for entry in entries:
            formatted.append(fmt.format(entry))
        result = {
            'command': 'bib',
            'action': 'import',
            'source': args.import_file,
            'total_entries': len(entries),
            'formatted': formatted,
        }
        msg = '\n'.join(f"[{i+1}] {r}" for i, r in enumerate(formatted))
        result['message'] = f"Imported {len(entries)} entries from {args.import_file}:\n{msg}"
        _emit(result, args)
    elif args.export_input:
        doc = docx_utils.open_docx(args.export_input)
        _, ref_entries = docx_utils.extract_reference_section(doc)
        refs = []
        for _, para, num in ref_entries:
            parsed = parse_ref(para.text.strip())
            parsed['key'] = f"ref{num}"
            refs.append(parsed)
        bib_output = export_bibtex(refs)
        if args.export_output:
            with open(args.export_output, 'w', encoding='utf-8') as f:
                f.write(bib_output)
        result = {
            'command': 'bib',
            'action': 'export',
            'source': args.export_input,
            'total_entries': len(refs),
            'bibtex': bib_output,
        }
        result['message'] = bib_output if not args.json else f"Exported {len(refs)} entries"
        _emit(result, args)
    else:
        err = {'error': 'Specify --import or --export-input'}
        print(json.dumps(err) if args.json else err['error'])
        sys.exit(1)


def cmd_interactive(args):
    """Interactive reference management console."""
    print(f"\nAcademic Ref Inserter v{__version__} - Interactive Mode")
    print("=" * 50)
    print("Commands: analyze | fix | generate | bib | formats | validate | help | quit")
    print("")

    while True:
        try:
            cmd_line = input("ref> ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\nGoodbye!")
            break

        if not cmd_line:
            continue
        if cmd_line.lower() in ('quit', 'exit', 'q'):
            print("Goodbye!")
            break
        if cmd_line.lower() in ('help', 'h', '?'):
            _print_interactive_help()
            continue
        if cmd_line.lower() in ('formats', 'fmts'):
            _print_formats()
            continue

        _handle_interactive_command(cmd_line)


def _print_interactive_help():
    print("""
Interactive Commands:
  analyze <file.docx>        Analyze citation structure
  check-refs <file.docx>     Auto-detect references and verify citation status
                               Options: --report FILE (export report to file)
  auto-find <file.docx>      Auto-search references and insert into document
                               Options: --query TITLE  --max-results N  --no-interactive
  fix <file.docx> <format>   Full pipeline (analysis -> reformat -> reorder -> hyperlink -> validate)
  validate <file.docx> [fmt] Validate references
  generate doi <DOI> <fmt>   Generate reference from DOI
  generate search <title> <fmt>  Search by title
  bib import <file.bib> <fmt>    Import from BibTeX
  bib export <file.docx> [out]  Export to BibTeX
  formats                    List available formats
  help                       Show this help
  quit                       Exit

Examples:
  ref> analyze paper.docx
  ref> check-refs paper.docx
  ref> auto-find paper.docx
  ref> auto-find paper.docx --query "time series forecasting"
  ref> fix paper.docx gbt7714
  ref> generate doi 10.1038/nature14539 ieee
  ref> formats
""")


def _print_formats():
    print("\nAvailable formats:")
    for key, fmt in FORMATS.items():
        print(f"  {key:12} {fmt.name:20} ordering: {fmt.ordering}")
    print("")


def _handle_interactive_command(cmd_line):
    parts = cmd_line.split()
    cmd = parts[0].lower()

    if cmd in ('check-refs', 'check') and len(parts) >= 2:
        class A:
            pass
        a = A()
        a.input = parts[1]
        a.output = None
        a.json = False
        a.format = 'gbt7714'
        try:
            cmd_check_refs(a)
        except Exception as e:
            print(f"Error: {e}")
    elif cmd in ('auto-find', 'autofind') and len(parts) >= 2:
        class A:
            pass
        a = A()
        a.input = parts[1]
        a.output = None
        a.json = False
        a.format = 'gbt7714'
        a.query = None
        a.max_results = 10
        a.no_interactive = False
        # Parse optional flags
        if '--query' in parts:
            qi = parts.index('--query')
            if qi + 1 < len(parts):
                a.query = parts[qi + 1]
        if '--max-results' in parts:
            mi = parts.index('--max-results')
            if mi + 1 < len(parts):
                try:
                    a.max_results = int(parts[mi + 1])
                except ValueError:
                    pass
        if '--no-interactive' in parts:
            a.no_interactive = True
        try:
            cmd_auto_find(a)
        except Exception as e:
            print(f"Error: {e}")
    elif cmd == 'analyze' and len(parts) >= 2:
        class A:
            pass
        a = A()
        a.input = parts[1]
        a.output = None
        a.json = False
        a.format = 'gbt7714'
        try:
            cmd_analyze(a)
        except Exception as e:
            print(f"Error: {e}")
    elif cmd == 'fix' and len(parts) >= 3:
        class A:
            pass
        a = A()
        a.input = parts[1]
        a.output = None
        a.json = False
        a.format = parts[2]
        try:
            cmd_fix(a)
        except Exception as e:
            print(f"Error: {e}")
    elif cmd == 'validate' and len(parts) >= 2:
        class A:
            pass
        a = A()
        a.input = parts[1]
        a.output = None
        a.json = False
        a.format = parts[2] if len(parts) >= 3 else 'gbt7714'
        try:
            cmd_validate(a)
        except Exception as e:
            print(f"Error: {e}")
    elif cmd == 'generate' and len(parts) >= 4:
        if parts[1] == 'doi':
            class A:
                pass
            a = A()
            a.doi = parts[2]
            a.format = parts[3]
            a.search = None
            a.json = False
            try:
                cmd_generate(a)
            except Exception as e:
                print(f"Error: {e}")
        elif parts[1] == 'search':
            class A:
                pass
            a = A()
            a.doi = None
            a.search = ' '.join(parts[2:])
            a.format = 'gbt7714'
            a.json = False
            try:
                cmd_generate(a)
            except Exception as e:
                print(f"Error: {e}")
    elif cmd == 'bib' and len(parts) >= 3:
        class A:
            pass
        a = A()
        a.json = False
        a.format = 'gbt7714'
        if parts[1] == 'import' and len(parts) >= 4:
            a.import_file = parts[2]
            a.format = parts[3]
            a.export_input = None
            a.export_output = None
            try:
                cmd_bib(a)
            except Exception as e:
                print(f"Error: {e}")
        elif parts[1] == 'export' and len(parts) >= 3:
            a.import_file = None
            a.export_input = parts[2]
            a.export_output = parts[3] if len(parts) >= 4 else None
            try:
                cmd_bib(a)
            except Exception as e:
                print(f"Error: {e}")
    else:
        print(f"Unknown command: {cmd}. Type 'help' for available commands.")


def main():
    parser = argparse.ArgumentParser(
        description='Academic Reference Inserter - format, reorder, and hyperlink references'
    )
    parser.add_argument('--version', '-V', action='version', version=f'%(prog)s {__version__}')
    parser.add_argument('--input', '-i', required=True, help='Input .docx file')
    parser.add_argument('--output', '-o', help='Output .docx file (default: overwrite input, safe with --output)')
    parser.add_argument('--format', '-f', choices=list(FORMATS.keys()), help='Citation format')
    parser.add_argument('--json', action='store_true', help='Output JSON for AI agent consumption')

    sub = parser.add_subparsers(dest='command', required=True)
    sub.add_parser('analyze', help='Analyze citation structure')
    p_auto = sub.add_parser('auto-find', help='Auto-search references and insert into document')
    p_auto.add_argument('--query', '-q', help='Search query (paper title or keywords). Auto-detected from document if omitted.')
    p_auto.add_argument('--max-results', '-m', type=int, default=10,
                        help='Maximum search results to return (default: 10)')
    p_auto.add_argument('--no-interactive', action='store_true',
                        help='Skip interactive selection, insert all results')
    p_check = sub.add_parser('check-refs', help='Auto-detect references and check citation status')
    p_check.add_argument('--report', '-r', help='Export report to file (.txt or .md)')
    sub.add_parser('reformat', help='Reformat references to target style')
    sub.add_parser('reorder', help='Reorder references to match citation order')
    sub.add_parser('hyperlink', help='Add cross-reference hyperlinks')
    sub.add_parser('validate', help='Validate references')
    sub.add_parser('fix', help='Full pipeline: all steps')
    sub_gen = sub.add_parser('generate', help='Generate reference from DOI or title search')
    sub_gen.add_argument('--doi', help='DOI string to look up')
    sub_gen.add_argument('--search', help='Search by title')
    sub_bib = sub.add_parser('bib', help='Import/export BibTeX')
    sub_bib.add_argument('--import', dest='import_file', help='Import from .bib file')
    sub_bib.add_argument('--export-input', help='Export from .docx to BibTeX')
    sub_bib.add_argument('--export-output', help='Output .bib file path')
    sub.add_parser('interactive', help='Interactive reference management console')

    args = parser.parse_args()

    if args.command in ('generate', 'bib', 'interactive'):
        try:
            commands = {
                'generate': cmd_generate,
                'bib': cmd_bib,
                'interactive': cmd_interactive,
            }
            commands[args.command](args)
        except Exception as e:
            err = {'error': str(e)}
            print(json.dumps(err, ensure_ascii=False) if args.json else str(e))
            sys.exit(1)
        return

    if not os.path.exists(args.input):
        err = {'error': f'File not found: {args.input}'}
        print(json.dumps(err) if args.json else err['error'])
        sys.exit(1)

    if not args.input.lower().endswith('.docx'):
        err = {'error': f'File must be .docx format: {args.input}'}
        print(json.dumps(err) if args.json else err['error'])
        sys.exit(1)

    try:
        commands = {
            'analyze': cmd_analyze,
            'check-refs': cmd_check_refs,
            'auto-find': cmd_auto_find,
            'reformat': cmd_reformat,
            'reorder': cmd_reorder,
            'hyperlink': cmd_hyperlink,
            'validate': cmd_validate,
            'fix': cmd_fix,
            'generate': cmd_generate,
            'bib': cmd_bib,
        }
        commands[args.command](args)
    except ValueError as e:
        err = {'error': str(e)}
        print(json.dumps(err, ensure_ascii=False) if args.json else str(e))
        sys.exit(1)
    except Exception as e:
        err = {
            'error': str(e),
            'hint': 'Ensure the file is a valid .docx document and try again.'
        }
        print(json.dumps(err, ensure_ascii=False) if args.json else str(e))
        sys.exit(1)


if __name__ == '__main__':
    main()
