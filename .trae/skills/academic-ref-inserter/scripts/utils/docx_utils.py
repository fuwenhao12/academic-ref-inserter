"""Word document (.docx) manipulation utilities.

Handles reading, writing, bookmark insertion, and hyperlink creation.
"""

import re
import shutil
from datetime import datetime
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def open_docx(path: str) -> Document:
    """Open a .docx file and return Document object."""
    return Document(path)


def backup_docx(path: str) -> str:
    """Create a timestamped backup of the docx file. Returns backup path."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = path.replace('.docx', f'_backup_refs_{timestamp}.docx')
    shutil.copy2(path, backup_path)
    return backup_path


def extract_all_text(doc: Document) -> list:
    """Extract all paragraph texts as a list of strings."""
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def extract_reference_section(doc: Document) -> tuple:
    """Extract the reference list section.

    Returns:
        (ref_title_para_index, ref_entries) where ref_entries is a list of
        (paragraph_index, paragraph_object, ref_number)
    """
    ref_title_idx = None
    ref_entries = []
    non_ref_count = 0
    max_blanks = 3

    section_headers = (
        '参考文献', 'References', 'REFERENCES', 'BIBLIOGRAPHY',
        '附录', 'Appendix', '致谢', 'Acknowledgments', 'Acknowledgements',
    )

    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()

        if ref_title_idx is None:
            if t in ('参考文献', 'References', 'REFERENCES', 'BIBLIOGRAPHY'):
                ref_title_idx = i
            continue

        # Stop if we hit another section
        if t in ('附录', 'Appendix', '致谢', 'Acknowledgments', 'Acknowledgements',
                 'Abstract', '摘要', 'Introduction', '引言', 'Conclusion', '结论',
                 '图表索引', 'List of Figures', 'List of Tables', 'Appendix'):
            break

        if re.match(r'^\[\d+\]', t):
            non_ref_count = 0
            num = int(re.match(r'\[(\d+)\]', t).group(1))
            ref_entries.append((i, para, num))
        elif not t:
            non_ref_count += 1
            if non_ref_count > max_blanks:
                break
        else:
            non_ref_count += 1
            if non_ref_count > 2:
                break

    return ref_title_idx, ref_entries


def extract_citations_from_text(doc: Document) -> list:
    """Extract all citation markers from body text (not reference section).

    Supports [N], [N-M], [N,M], [N,M-O] formats.
    Returns list of citation numbers in order of first appearance.
    """
    first_occurrence = []
    seen_nums = set()

    ref_end = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() in ('参考文献', 'References', 'REFERENCES', 'BIBLIOGRAPHY'):
            ref_end = i
            break

    for i, para in enumerate(doc.paragraphs):
        if ref_end is not None and i >= ref_end:
            break
        text = para.text.strip()
        if not text:
            continue
        for match in re.finditer(r'\[([0-9,\-\s]+)\]', text):
            content = match.group(1)
            for num in _parse_citation_range(content):
                if num not in seen_nums:
                    seen_nums.add(num)
                    first_occurrence.append(num)

    return first_occurrence


def _parse_citation_range(content: str) -> list:
    """Parse citation content like '1-3,5,7-9' into flat list of ints."""
    nums = []
    for part in content.split(','):
        part = part.strip()
        if '-' in part:
            a, _, b = part.partition('-')
            try:
                for n in range(int(a.strip()), int(b.strip()) + 1):
                    nums.append(n)
            except ValueError:
                pass
        else:
            try:
                nums.append(int(part.strip()))
            except ValueError:
                pass
    return nums


def extract_citation_context(doc: Document, num: int) -> str:
    """Find the context around a citation marker [num] in body text."""
    for para in doc.paragraphs:
        text = para.text
        pattern = f'[{num}]'
        if pattern in text:
            idx = text.find(pattern)
            start = max(0, idx - 20)
            end = min(len(text), idx + len(pattern) + 30)
            return text[start:end]
    return ''


def add_bookmark_to_paragraph(para, bookmark_name: str) -> bool:
    """Add a bookmark at the start of a paragraph."""
    try:
        bookmark_id = str(abs(hash(bookmark_name)) % 1000000)
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), bookmark_id)
        bookmark_start.set(qn('w:name'), bookmark_name)

        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), bookmark_id)

        first_run = para.runs[0] if para.runs else None
        if first_run:
            first_run._r.addprevious(bookmark_start)
            para._p.append(bookmark_end)
            return True
        return False
    except Exception:
        return False


def make_hyperlink_element(ref_text: str, bookmark_name: str,
                           superscript: bool = True,
                           color: str = None) -> OxmlElement:
    """Create a w:hyperlink XML element pointing to a bookmark.

    Args:
        ref_text: The citation text to display (e.g. '[1]')
        bookmark_name: Target bookmark name
        superscript: Whether to apply superscript styling
        color: Font color hex (e.g. '000000'), None = inherit from document
    """
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if superscript:
        vAlign = OxmlElement('w:vertAlign')
        vAlign.set(qn('w:val'), 'superscript')
        rPr.append(vAlign)

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = ref_text
    r.append(t)
    hyperlink.append(r)
    return hyperlink


def replace_citation_with_hyperlink(para, ref_num: int) -> bool:
    """Replace ALL occurrences of [ref_num] in a paragraph with hyperlinks.

    Only processes runs that are direct children of <w:p> (skips hyperlink runs).
    Returns True if at least one replacement succeeded.
    """
    ref_text = f'[{ref_num}]'
    bookmark_name = f"Ref_{ref_num}"
    replaced_any = False

    # Only iterate through direct <w:r> children of <w:p>, skipping hyperlink runs
    p = para._p
    while True:
        direct_runs = [r for r in p if r.tag == qn('w:r')]
        found_run = None
        found_offset = -1
        char_count = 0

        for r in direct_runs:
            run_text = r.text or ''
            idx = run_text.find(ref_text)
            if idx >= 0:
                found_run = r
                found_offset = idx
                break
            char_count += len(run_text)

        if found_run is None:
            break

        before = found_run.text[:found_offset]
        after = found_run.text[found_offset + len(ref_text):]

        hyperlink_elem = make_hyperlink_element(ref_text, bookmark_name)

        if before:
            found_run.text = before
            found_run.addnext(hyperlink_elem)
        else:
            found_run.text = ''
            found_run.addprevious(hyperlink_elem)

        if after:
            new_run = OxmlElement('w:r')
            rPr = None
            for child in found_run:
                if child.tag == qn('w:rPr'):
                    rPr = deepcopy(child)
                    break
            if rPr:
                new_run.append(deepcopy(rPr))
            new_t = OxmlElement('w:t')
            new_t.text = after
            new_run.append(new_t)
            if before:
                hyperlink_elem.addnext(new_run)
            else:
                found_run.addnext(new_run)

        replaced_any = True

    return replaced_any


def insert_text_at_run(para, search_text: str, insert_text: str) -> bool:
    """Insert text after the first occurrence of search_text in any run."""
    for run in para.runs:
        if search_text in run.text:
            idx = run.text.index(search_text) + len(search_text)
            run.text = run.text[:idx] + insert_text + run.text[idx:]
            return True
    return False


def get_reference_text(ref_entries: list, ref_num: int) -> str:
    """Get the full text of a reference entry by its number."""
    for idx, para, num in ref_entries:
        if num == ref_num:
            return para.text.strip()
    return ''


def update_reference_text(para, new_text: str) -> bool:
    """Update the text of a reference paragraph."""
    if not para.runs:
        return False
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ''
    return True


def dedup_adjacent_citations(para) -> int:
    """Remove duplicate adjacent citations like [1][1][1] -> [1] in a paragraph.

    Uses iterative regex replacement across the paragraph's full text,
    then writes the cleaned text back into the first run.
    Returns number of duplicate markers removed.
    """
    full_text = para.text
    prev_text = None
    while prev_text != full_text:
        prev_text = full_text
        full_text = re.sub(r'\[(\d+)\]\s*\[(\1)\]', r'[\1]', full_text)

    removed_count = (len(para.text) - len(full_text)) // 2  # approx count
    if full_text == para.text:
        return 0

    if para.runs:
        para.runs[0].text = full_text
        for r in para.runs[1:]:
            r.text = ''
    return removed_count


def find_reference_boundary_robust(doc: Document) -> int:
    """Robustly find where the reference section starts.

    Tries standard section headers first, then falls back to
    detecting a dense block of [N]-prefixed paragraphs.
    Returns paragraph index (0-based), or None if not found.
    """
    # Attempt 1: Standard section headers
    section_headers = (
        '参考文献', 'References', 'REFERENCES', 'BIBLIOGRAPHY',
    )
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if t in section_headers:
            return i

    # Attempt 2: Detect dense [N]-prefixed block
    # Scan windows of paragraphs; if 3+ out of 5 consecutive paragraphs
    # match [N] pattern, consider that the start of the reference section.
    for i in range(len(doc.paragraphs) - 5):
        match_count = 0
        for j in range(5):
            t = doc.paragraphs[i + j].text.strip()
            if re.match(r'^\[\d+\]', t):
                match_count += 1
        if match_count >= 3:
            # Walk back to find the actual start
            start = i
            while start > 0:
                prev_text = doc.paragraphs[start - 1].text.strip()
                if re.match(r'^\[\d+\]', prev_text):
                    start -= 1
                else:
                    break
            return start

    return None


def find_citation_occurrences(doc: Document, ref_num: int,
                              context_chars: int = 30) -> list:
    """Find all paragraphs where [ref_num] is cited.

    Args:
        doc: Document object
        ref_num: Reference number to search for
        context_chars: Chars of context to include around the citation

    Returns:
        List of dicts: {'para_index': N, 'snippet': '...text...'}
        Empty list if reference is not cited.
    """
    ref_end = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() in ('参考文献', 'References', 'REFERENCES', 'BIBLIOGRAPHY'):
            ref_end = i
            break
    if ref_end is None:
        ref_end = len(doc.paragraphs)

    occurrences = []
    pattern = f'[{ref_num}]'

    for i, para in enumerate(doc.paragraphs):
        if i >= ref_end:
            break
        text = para.text
        if pattern in text:
            idx = text.find(pattern)
            start = max(0, idx - context_chars)
            end = min(len(text), idx + len(pattern) + context_chars)
            snippet = text[start:end]
            occurrences.append({
                'para_index': i,
                'snippet': snippet,
            })

    return occurrences


def extract_citation_content(content: str) -> list:
    """Parse citation content like '1-3,5,7-9' into flat list of ints."""
    nums = []
    for part in content.split(','):
        part = part.strip()
        if '-' in part:
            a, _, b = part.partition('-')
            try:
                for n in range(int(a.strip()), int(b.strip()) + 1):
                    nums.append(n)
            except ValueError:
                pass
        else:
            try:
                nums.append(int(part.strip()))
            except ValueError:
                pass
    return nums
